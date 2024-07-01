import os
from docx import Document
from pathlib import Path

# 使用 os 模块递归查找所有文件
def find_all_files_os(folder_path):
    all_files = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            all_files.append(os.path.join(root, file))
    return all_files

# 使用 pathlib 模块递归查找所有文件
def find_all_files_pathlib(folder_path):
    path = Path(folder_path)
    return [str(file) for file in path.rglob('*') if file.is_file()]

# 从所有文件中提取内容并写入到一个 .docx 文件中
def extract_text_files_to_docx(folder_path, output_docx, use_pathlib=True):
    document = Document()

    if use_pathlib:
        all_files = find_all_files_pathlib(folder_path)
    else:
        all_files = find_all_files_os(folder_path)

    for file_path in all_files:
        if file_path.endswith(".vue"):
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                document.add_heading(Path(file_path).name, level=1)
                document.add_paragraph(content)

    document.save(output_docx)
    print(f"内容已提取并保存到 {output_docx}")

# 指定文件夹路径和输出的 docx 文件名
folder_path = 'pages'
output_docx = 'output.docx'

# 执行提取和写入操作
extract_text_files_to_docx(folder_path, output_docx, use_pathlib=False)